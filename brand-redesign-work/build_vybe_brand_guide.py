from pathlib import Path
import math
import os
import subprocess

import numpy as np
from PIL import Image, ImageDraw, ImageFont, ImageFilter
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


ROOT = Path("/Users/ayorindejohn/Documents/VYBE_CENTRAL")
WORK = ROOT / "brand-redesign-work"
SLIDES = WORK / "slides"
ASSETS = WORK / "assets"
OUT_DOCX = ROOT / "Vybe_Central_Brand_Guidelines_2026.docx"
OUT_PDF_DIRECT = ROOT / "Vybe_Central_Brand_Guidelines_2026_preview.pdf"

LOGO_SRC = Path("/Users/ayorindejohn/Downloads/Professional_recording_studio_logo_for_202606101100.jpeg")
GEN_DIR = Path("/Users/ayorindejohn/.codex/generated_images/019eb141-b5da-7df0-b2e2-61885e340f15")

FONT_DIR = Path(
    "/Users/ayorindejohn/.cache/codex-runtimes/codex-primary-runtime/dependencies/native/"
    "libreoffice-headless/libreoffice/LibreOfficeDev.app/Contents/Resources/fonts/truetype"
)
PYTHON = Path("/Users/ayorindejohn/.cache/codex-runtimes/codex-primary-runtime/dependencies/python/bin/python3")

W, H = 3200, 2500

BLACK = "#080808"
INK = "#121212"
CHARCOAL = "#1B1715"
CREAM = "#F4EAD8"
CREAM_2 = "#FFF6E8"
ORANGE = "#F08A24"
ORANGE_2 = "#E85A24"
GOLD = "#F5B43A"
BLUE = "#67C5E7"
BLUE_DARK = "#1E6C88"
WHITE = "#FFFFFF"
MUTED = "#A89C8C"
RED = "#E6462B"


def rgb(hex_color):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def font(name, size):
    return ImageFont.truetype(str(FONT_DIR / name), size)


DISPLAY = "LiberationSansNarrow-Bold.ttf"
DISPLAY_REG = "LiberationSansNarrow-Regular.ttf"
BODY = "Rubik-Regular.ttf"
BODY_BOLD = "Rubik-Bold.ttf"
BODY_ITALIC = "Rubik-Italic.ttf"
MONO = "DejaVuSansMono.ttf"


def text_size(draw, text, fnt):
    box = draw.textbbox((0, 0), text, font=fnt)
    return box[2] - box[0], box[3] - box[1]


def wrap_text(draw, text, fnt, max_width):
    lines = []
    for raw in text.split("\n"):
        words = raw.split()
        if not words:
            lines.append("")
            continue
        line = words[0]
        for word in words[1:]:
            test = f"{line} {word}"
            if text_size(draw, test, fnt)[0] <= max_width:
                line = test
            else:
                lines.append(line)
                line = word
        lines.append(line)
    return lines


def draw_wrapped(draw, xy, text, fnt, fill, max_width, leading=1.22, anchor=None):
    x, y = xy
    lines = wrap_text(draw, text, fnt, max_width)
    step = int(fnt.size * leading)
    for line in lines:
        if anchor == "center":
            tw, _ = text_size(draw, line, fnt)
            draw.text((x - tw / 2, y), line, font=fnt, fill=fill)
        else:
            draw.text((x, y), line, font=fnt, fill=fill)
        y += step
    return y


def draw_title(draw, x, y, lines, colors=None, size=210, max_width=None):
    fnt = font(DISPLAY, size)
    colors = colors or [WHITE] * len(lines)
    gap = int(size * 0.82)
    for i, line in enumerate(lines):
        if max_width:
            while text_size(draw, line, fnt)[0] > max_width and size > 80:
                size -= 4
                fnt = font(DISPLAY, size)
                gap = int(size * 0.82)
        draw.text((x, y + i * gap), line, font=fnt, fill=colors[i % len(colors)])
    return y + len(lines) * gap


def gradient_bg(left=BLACK, right=INK):
    a = np.array(rgb(left), dtype=np.float32)
    b = np.array(rgb(right), dtype=np.float32)
    grad = np.linspace(0, 1, W, dtype=np.float32)
    row = (a * (1 - grad[:, None]) + b * grad[:, None]).astype(np.uint8)
    arr = np.repeat(row[None, :, :], H, axis=0)
    return Image.fromarray(arr, "RGB").convert("RGBA")


def radial_overlay(img, color, center, radius, strength=0.55):
    overlay = Image.new("RGBA", (W, H), rgb(color) + (0,))
    arr = np.zeros((H, W), dtype=np.float32)
    yy, xx = np.ogrid[:H, :W]
    dist = np.sqrt((xx - center[0]) ** 2 + (yy - center[1]) ** 2)
    arr = np.clip(1 - dist / radius, 0, 1) * int(255 * strength)
    alpha = Image.fromarray(arr.astype(np.uint8), "L").filter(ImageFilter.GaussianBlur(36))
    overlay.putalpha(alpha)
    return Image.alpha_composite(img, overlay)


def draw_grid(draw, color="#24201d", alpha=120):
    col = rgb(color) + (alpha,)
    for x in range(430, W, 430):
        draw.line((x, 0, x, H), fill=col, width=2)
    for y in range(360, H, 360):
        draw.line((0, y, W, y), fill=col, width=1)


def draw_wave(draw, x, y, w, h, color=ORANGE, steps=64):
    pts = []
    for i in range(steps + 1):
        t = i / steps
        amp = math.sin(t * math.pi) * h * 0.42
        yy = y + h / 2 + math.sin(t * math.pi * 8) * amp
        pts.append((x + t * w, yy))
    draw.line(pts, fill=rgb(color) + (220,), width=8, joint="curve")
    for i in range(0, steps + 1, 4):
        px = x + i / steps * w
        draw.line((px, y + h / 2 - 18, px, y + h / 2 + 18), fill=rgb(CREAM) + (90,), width=3)


def pill(draw, xy, text, fill, fg=BLACK, pad_x=30, pad_y=12, fnt=None):
    fnt = fnt or font(BODY_BOLD, 38)
    x, y = xy
    tw, th = text_size(draw, text, fnt)
    box = (x, y, x + tw + pad_x * 2, y + th + pad_y * 2)
    draw.rounded_rectangle(box, radius=42, fill=rgb(fill) + (255,))
    draw.text((x + pad_x, y + pad_y - 2), text, font=fnt, fill=rgb(fg) + (255,))
    return box[2], box[3]


def card(draw, box, fill, outline=None, radius=34, width=3):
    draw.rounded_rectangle(box, radius=radius, fill=rgb(fill) + (255,))
    if outline:
        draw.rounded_rectangle(box, radius=radius, outline=rgb(outline) + (255,), width=width)


def cover_crop(path, size, focus=(0.5, 0.5)):
    img = Image.open(path).convert("RGB")
    sw, sh = img.size
    tw, th = size
    scale = max(tw / sw, th / sh)
    nw, nh = int(sw * scale), int(sh * scale)
    img = img.resize((nw, nh), Image.Resampling.LANCZOS)
    left = int((nw - tw) * focus[0])
    top = int((nh - th) * focus[1])
    return img.crop((left, top, left + tw, top + th)).convert("RGBA")


def tint_image(img, color=BLACK, alpha=130):
    layer = Image.new("RGBA", img.size, rgb(color) + (alpha,))
    return Image.alpha_composite(img.convert("RGBA"), layer)


def make_logo_cutout():
    ASSETS.mkdir(parents=True, exist_ok=True)
    out = ASSETS / "vybe-logo-cutout.png"
    img = Image.open(LOGO_SRC).convert("RGBA")
    arr = np.array(img.convert("RGB"))
    mean = arr.mean(axis=2)
    sat = arr.max(axis=2) - arr.min(axis=2)
    obj = (mean < 245) & ((sat > 10) | (mean < 220))
    ys, xs = np.where(obj)
    left, top, right, bottom = xs.min(), ys.min(), xs.max(), ys.max()
    cx = (left + right) / 2
    cy = (top + bottom) / 2
    radius = max(right - left, bottom - top) / 2 + 34
    yy, xx = np.ogrid[: img.height, : img.width]
    dist = np.sqrt((xx - cx) ** 2 + (yy - cy) ** 2)
    alpha = np.clip((radius - dist) / 24, 0, 1) * 255
    alpha = alpha.astype(np.uint8)
    out_arr = np.array(img)
    out_arr[:, :, 3] = alpha
    cut = Image.fromarray(out_arr, "RGBA")
    pad = 90
    crop = (
        max(0, int(cx - radius - pad)),
        max(0, int(cy - radius - pad)),
        min(img.width, int(cx + radius + pad)),
        min(img.height, int(cy + radius + pad)),
    )
    cut = cut.crop(crop)
    cut.save(out)
    return out


def paste_logo(page, logo_path, box, shadow=True):
    logo = Image.open(logo_path).convert("RGBA")
    logo.thumbnail((box[2] - box[0], box[3] - box[1]), Image.Resampling.LANCZOS)
    x = box[0] + ((box[2] - box[0]) - logo.width) // 2
    y = box[1] + ((box[3] - box[1]) - logo.height) // 2
    if shadow:
        sh = Image.new("RGBA", logo.size, (0, 0, 0, 0))
        sh.putalpha(logo.getchannel("A").filter(ImageFilter.GaussianBlur(22)))
        page.alpha_composite(sh, (x + 28, y + 34))
    page.alpha_composite(logo, (x, y))


def page_base(bg=BLACK, grid=True):
    img = gradient_bg(bg, bg if bg != BLACK else "#120f0d")
    img = radial_overlay(img, ORANGE, (W + 120, 180), 1400, 0.28)
    d = ImageDraw.Draw(img, "RGBA")
    if grid:
        draw_grid(d)
    return img, d


def page_number(draw, n, dark=False):
    f = font(MONO, 30)
    fill = rgb(INK if dark else CREAM) + (150,)
    draw.text((W - 255, 118), f"2026 / {n:02d}", font=f, fill=fill)


def top_mark(draw, logo_path=None, dark=False):
    f = font(BODY_BOLD, 46)
    fill = rgb(INK if dark else CREAM) + (230,)
    draw.text((142, 124), "VYBE CENTRAL", font=f, fill=fill)
    draw_wave(draw, W - 545, 96, 170, 92, color=ORANGE if not dark else INK)


def save_slide(img, idx):
    SLIDES.mkdir(parents=True, exist_ok=True)
    p = SLIDES / f"page-{idx:02d}.png"
    img.convert("RGB").save(p, quality=95)
    return p


def add_bullets(draw, x, y, items, width, fnt=None, color=CREAM, bullet_color=ORANGE, gap=32):
    fnt = fnt or font(BODY, 52)
    for item in items:
        draw.ellipse((x, y + 18, x + 18, y + 36), fill=rgb(bullet_color) + (255,))
        y = draw_wrapped(draw, (x + 42, y), item, fnt, rgb(color) + (235,), width - 42, 1.24)
        y += gap
    return y


def draw_swatches(draw, swatches, x, y, w, h):
    gap = 34
    col_w = (w - gap) // 2
    row_h = h // 4 - 18
    for i, (name, hexv, note) in enumerate(swatches):
        col = i % 2
        row = i // 2
        bx = x + col * (col_w + gap)
        by = y + row * (row_h + 36)
        card(draw, (bx, by, bx + col_w, by + row_h), hexv, outline="#000000" if hexv in [CREAM, CREAM_2, WHITE] else None, radius=20)
        label = font(BODY_BOLD, 34)
        body = font(BODY, 27)
        fg = BLACK if hexv in [CREAM, CREAM_2, WHITE, BLUE, GOLD, ORANGE] else CREAM
        draw.text((bx + 34, by + 26), name.upper(), font=label, fill=rgb(fg) + (255,))
        draw.text((bx + 34, by + 72), hexv, font=font(MONO, 28), fill=rgb(fg) + (230,))
        draw_wrapped(draw, (bx + 34, by + 116), note, body, rgb(fg) + (230,), col_w - 68, 1.16)


def build_pages():
    logo = make_logo_cutout()
    generated = sorted(GEN_DIR.glob("*.png")) if GEN_DIR.exists() else []
    hero = generated[0] if len(generated) > 0 else LOGO_SRC
    studio = generated[1] if len(generated) > 1 else LOGO_SRC
    artist = generated[2] if len(generated) > 2 else hero
    pages = []

    # 01 Cover
    img, d = page_base(BLACK)
    bg = cover_crop(hero, (1750, H), focus=(0.70, 0.50))
    bg = tint_image(bg, BLACK, 86)
    img.alpha_composite(bg, (W - 1750, 0))
    d.rectangle((W - 1750, 0, W, H), fill=rgb(BLACK) + (55,))
    top_mark(d)
    page_number(d, 1)
    pill(d, (154, 496), "UPDATED WEBSITE BRAND SYSTEM", ORANGE, BLACK)
    y = draw_title(d, 145, 690, ["BRAND", "GUIDE"], [CREAM, ORANGE], size=280, max_width=1340)
    d.text((150, y + 38), "Studio. Sessions. Community. Discovery.", font=font(BODY_BOLD, 64), fill=rgb(BLUE) + (255,))
    draw_wrapped(
        d,
        (152, y + 130),
        "A refreshed visual and messaging system for VYBE Central as a recording studio, live performance platform, creative community, and launchpad for Sierra Leonean and African talent.",
        font(BODY, 48),
        rgb(CREAM) + (220,),
        1130,
        1.28,
    )
    paste_logo(img, logo, (2210, 835, 3020, 1645), shadow=True)
    draw_wave(d, 150, H - 302, 530, 128, BLUE)
    d.text((154, H - 148), "VYBE CENTRAL / 2026", font=font(MONO, 34), fill=rgb(CREAM) + (155,))
    pages.append(save_slide(img, 1))

    # 02 Brand overview
    img = gradient_bg(CREAM, CREAM_2)
    d = ImageDraw.Draw(img, "RGBA")
    d.rectangle((0, 0, 1180, H), fill=rgb(INK) + (255,))
    side = cover_crop(artist, (1160, 1020), focus=(0.55, 0.45))
    img.alpha_composite(side, (102, H - 1160))
    d.rectangle((102, H - 1160, 1262, H - 140), fill=rgb(BLACK) + (35,))
    d.rectangle((102, H - 1160, 1262, H - 140), outline=rgb(ORANGE) + (70,), width=3)
    draw_wave(d, 220, H - 880, 820, 210, BLUE)
    d.text((225, H - 585), "STUDIO", font=font(DISPLAY, 100), fill=rgb(CREAM) + (255,))
    d.text((225, H - 465), "SESSIONS", font=font(DISPLAY, 100), fill=rgb(ORANGE) + (255,))
    d.text((225, H - 345), "COMMUNITY", font=font(DISPLAY, 100), fill=rgb(CREAM) + (255,))
    d.text((225, H - 225), "PLATFORM", font=font(DISPLAY, 100), fill=rgb(BLUE) + (255,))
    top_mark(d, dark=True)
    page_number(d, 2, dark=True)
    d.text((1330, 248), "BRAND OVERVIEW", font=font(DISPLAY, 168), fill=rgb(INK) + (255,))
    pill(d, (1336, 435), "WHAT VYBE CENTRAL IS", ORANGE, BLACK, fnt=font(BODY_BOLD, 34))
    draw_wrapped(
        d,
        (1335, 535),
        "VYBE Central is a creative ecosystem where artists, producers, storytellers, and creators come together to create, collaborate, and be heard.",
        font(BODY_BOLD, 62),
        rgb(INK) + (255,),
        1500,
        1.20,
    )
    sections = [
        ("MISSION", "Amplify African voices through studio production, live performance sessions, storytelling, and creative community."),
        ("VISION", "Become a platform where Sierra Leonean and African talent is discovered, supported, and shared with the world."),
        ("PROMISE", "Professional creation space, culturally rooted storytelling, and visibility beyond the session."),
    ]
    y = 1000
    for label, text in sections:
        d.text((1350, y), label, font=font(DISPLAY, 64), fill=rgb(ORANGE) + (255,))
        draw_wrapped(d, (1352, y + 70), text, font(BODY, 40), rgb(INK) + (225,), 1210, 1.24)
        y += 265
    d.text((140, 225), "ESSENCE", font=font(DISPLAY, 122), fill=rgb(CREAM) + (255,))
    add_bullets(
        d,
        145,
        415,
        ["Creative home base", "Performance-first content", "Community over gatekeeping", "African talent discovery"],
        810,
        fnt=font(BODY, 42),
    )
    pages.append(save_slide(img, 2))

    # 03 Positioning
    img, d = page_base(BLACK)
    mic = cover_crop(studio, (1060, 1490), focus=(0.50, 0.45))
    mic = tint_image(mic, BLACK, 30)
    img.alpha_composite(mic, (2030, 550))
    top_mark(d)
    page_number(d, 3)
    d.text((145, 320), "WHY VYBE", font=font(DISPLAY, 190), fill=rgb(CREAM) + (255,))
    d.text((145, 500), "CENTRAL?", font=font(DISPLAY, 190), fill=rgb(ORANGE) + (255,))
    draw_wrapped(
        d,
        (152, 748),
        "The brand should feel like a working studio and a cultural platform at the same time: precise, cinematic, social, and rooted.",
        font(BODY_BOLD, 60),
        rgb(CREAM) + (238,),
        1320,
        1.20,
    )
    pillars = [
        ("RECORDING STUDIO", "Music, voice-over, rehearsal, and artist development."),
        ("LIVE SESSIONS", "Performance content designed for discovery and sharing."),
        ("CREATIVE COMMUNITY", "A place for producers, storytellers, hosts, and creators to meet."),
        ("DISCOVER ARTISTS", "Profiles, bios, social links, and performance video for every session artist."),
    ]
    x0, y0 = 150, 1120
    for i, (label, body) in enumerate(pillars):
        x = x0 + (i % 2) * 840
        y = y0 + (i // 2) * 405
        card(d, (x, y, x + 730, y + 290), "#14100E", outline="#332A24", radius=22)
        d.text((x + 36, y + 34), f"0{i+1}", font=font(MONO, 44), fill=rgb(BLUE) + (255,))
        d.text((x + 36, y + 92), label, font=font(DISPLAY, 58), fill=rgb(ORANGE) + (255,))
        draw_wrapped(d, (x + 38, y + 160), body, font(BODY, 34), rgb(CREAM) + (210,), 620, 1.18)
    pages.append(save_slide(img, 3))

    # 04 Website system
    img = gradient_bg(BLUE, "#7ED2EF")
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d, dark=True)
    page_number(d, 4, dark=True)
    d.text((140, 278), "WEBSITE", font=font(DISPLAY, 180), fill=rgb(INK) + (255,))
    d.text((140, 444), "EXPERIENCE MAP", font=font(DISPLAY, 180), fill=rgb(INK) + (255,))
    d.text((154, 664), "The site should behave like a platform, not a brochure.", font=font(BODY_BOLD, 56), fill=rgb(INK) + (230,))
    flow = [
        ("01", "HERO", "Immediate signal: studio, artist, live-session energy."),
        ("02", "CREATE", "Recording studio, rehearsal, live sessions, voice-over."),
        ("03", "JOIN", "Capture email and WhatsApp with a tight community form."),
        ("04", "CALLS", "Rotating open calls for vocalists, musicians, hosts, producers."),
        ("05", "DISCOVER", "Artist profiles with photos, bios, social links, and video."),
    ]
    y = 900
    for num, head, body in flow:
        d.text((160, y + 18), num, font=font(DISPLAY, 76), fill=rgb(INK) + (255,))
        card(d, (315, y, 1760, y + 190), CREAM_2, outline="#111111", radius=16)
        d.text((360, y + 34), head, font=font(DISPLAY, 64), fill=rgb(ORANGE_2) + (255,))
        draw_wrapped(d, (650, y + 48), body, font(BODY, 34), rgb(INK) + (230,), 920, 1.15)
        y += 230
    shot = cover_crop(hero, (1030, 1110), focus=(0.60, 0.50))
    img.alpha_composite(shot, (2030, 760))
    d.rectangle((2030, 760, 3060, 1870), outline=rgb(INK) + (255,), width=10)
    draw_wrapped(
        d,
        (2050, 1930),
        "WEBSITE RULE: every section must lead to action, proof, or discovery.",
        font(BODY_BOLD, 34),
        rgb(INK) + (230,),
        910,
        1.12,
    )
    pages.append(save_slide(img, 4))

    # 05 Logo system
    img = gradient_bg(CREAM, CREAM_2)
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d, dark=True)
    page_number(d, 5, dark=True)
    d.text((140, 265), "LOGO SYSTEM", font=font(DISPLAY, 172), fill=rgb(INK) + (255,))
    draw_wrapped(
        d,
        (147, 460),
        "Use the circular studio mark as the main identity. The background has been removed so the logo can sit cleanly on dark, cream, orange, and blue surfaces.",
        font(BODY, 44),
        rgb(INK) + (220,),
        1220,
        1.22,
    )
    paste_logo(img, logo, (178, 690, 1290, 1802), shadow=True)
    d.ellipse((150, 662, 1320, 1832), outline=rgb(ORANGE) + (255,), width=8)
    variants = [(1490, 715, BLACK, CREAM), (2080, 715, CREAM, BLACK), (1490, 1285, ORANGE, BLACK), (2080, 1285, BLUE, BLACK)]
    for x, y, bgc, fg in variants:
        card(d, (x, y, x + 500, y + 420), bgc, outline="#222222" if bgc in [CREAM, BLUE] else None, radius=12)
        paste_logo(img, logo, (x + 80, y + 42, x + 420, y + 382), shadow=False)
    d.text((1490, 530), "APPROVED PLACEMENTS", font=font(DISPLAY, 72), fill=rgb(ORANGE_2) + (255,))
    rules = [
        "Keep clear space equal to the microphone width around the mark.",
        "Never place the logo on a busy image without a dark overlay.",
        "Use the full circular mark for hero, footer, and profile surfaces.",
        "Use simple text lockups only when space is too tight for the circle.",
    ]
    add_bullets(d, 1488, 1790, rules, 1220, fnt=font(BODY, 35), color=INK, bullet_color=ORANGE, gap=18)
    pages.append(save_slide(img, 5))

    # 06 Color system
    img = gradient_bg("#E4352C", "#F89B20")
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d)
    page_number(d, 6)
    d.text((140, 275), "COLOR SYSTEM", font=font(DISPLAY, 172), fill=rgb(WHITE) + (255,))
    draw_wrapped(
        d,
        (146, 462),
        "The website palette fuses the black and orange of the logo with the cream editorial feel and blue utility band from the visual inspiration.",
        font(BODY_BOLD, 50),
        rgb(WHITE) + (238,),
        1370,
        1.18,
    )
    swatches = [
        ("Studio Black", BLACK, "Primary surface. Use for hero, footer, nav, and performance sections."),
        ("Warm Cream", CREAM, "Editorial pages, artist cards, readable sections, and background space."),
        ("VYBE Orange", ORANGE, "Primary action color, key labels, hover states, and active motion."),
        ("Flame Red", RED, "Secondary heat for badges, urgency, calls for talent, and campaign moments."),
        ("Session Blue", BLUE, "Utility bands, community capture, form focus, and system accents."),
        ("Gold Hit", GOLD, "Small highlight only. Use for rhythm, not large backgrounds."),
        ("Charcoal", CHARCOAL, "Cards and panels on dark surfaces."),
        ("White", WHITE, "Logo contrast, captions, and crisp UI details."),
    ]
    draw_swatches(d, swatches, 1430, 520, 1540, 1640)
    d.text((150, 1990), "RATIO: 55% black / 25% cream / 12% orange-red / 8% blue-gold accents", font=font(BODY_BOLD, 42), fill=rgb(BLACK) + (225,))
    pages.append(save_slide(img, 6))

    # 07 Typography
    img, d = page_base(BLACK)
    top_mark(d)
    page_number(d, 7)
    d.text((140, 285), "TYPOGRAPHY", font=font(DISPLAY, 176), fill=rgb(ORANGE) + (255,))
    d.text((150, 520), "DISPLAY", font=font(DISPLAY, 92), fill=rgb(CREAM) + (255,))
    d.text((148, 635), "BEBAS NEUE / ANTON", font=font(DISPLAY, 142), fill=rgb(CREAM) + (255,))
    d.text((150, 790), "Fallback: Liberation Sans Narrow Bold", font=font(BODY, 44), fill=rgb(MUTED) + (230,))
    d.text((150, 1005), "BODY", font=font(DISPLAY, 92), fill=rgb(BLUE) + (255,))
    d.text((150, 1120), "Rubik or Inter for paragraphs, forms, captions, and UI.", font=font(BODY, 62), fill=rgb(CREAM) + (238,))
    d.text((150, 1295), "A B C D E F G H I J K L M", font=font(BODY_BOLD, 56), fill=rgb(WHITE) + (210,))
    d.text((150, 1380), "a b c d e f g h i j k l m", font=font(BODY, 54), fill=rgb(WHITE) + (185,))
    d.text((150, 1540), "STYLE NOTES", font=font(DISPLAY, 82), fill=rgb(ORANGE) + (255,))
    add_bullets(
        d,
        155,
        1650,
        [
            "Headlines are tight, uppercase, and cinematic.",
            "Body copy stays calm and readable.",
            "Accent words can use italic or script, but only once per section.",
            "No negative letter spacing. Let the type breathe.",
        ],
        1240,
        fnt=font(BODY, 39),
    )
    d.text((1780, 560), "WE CRAFT", font=font(DISPLAY, 178), fill=rgb(WHITE) + (255,))
    d.text((1780, 720), "LIVE SESSIONS", font=font(DISPLAY, 178), fill=rgb(ORANGE) + (255,))
    d.text((1782, 918), "THAT HELP ARTISTS", font=font(DISPLAY, 116), fill=rgb(CREAM) + (255,))
    d.text((1782, 1040), "BE DISCOVERED.", font=font(DISPLAY, 116), fill=rgb(BLUE) + (255,))
    pages.append(save_slide(img, 7))

    # 08 Visual language
    img = gradient_bg(CREAM, CREAM_2)
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d, dark=True)
    page_number(d, 8, dark=True)
    d.text((140, 248), "VISUAL LANGUAGE", font=font(DISPLAY, 154), fill=rgb(INK) + (255,))
    d.text((145, 425), "Cinematic, editorial, and rooted in real creative work.", font=font(BODY_BOLD, 52), fill=rgb(INK) + (230,))
    pic1 = cover_crop(hero, (880, 960), focus=(0.45, 0.5))
    pic2 = cover_crop(studio, (790, 780), focus=(0.50, 0.45))
    pic3 = cover_crop(artist, (720, 780), focus=(0.55, 0.45))
    d.rectangle((160, 690, 1040, 1650), fill=rgb(ORANGE_2) + (255,))
    img.alpha_composite(pic1, (215, 760))
    d.rectangle((1270, 615, 2060, 1395), fill=rgb(BLACK) + (255,))
    img.alpha_composite(pic2, (1325, 680))
    d.rectangle((2170, 950, 2890, 1730), fill=rgb(BLACK) + (255,))
    img.alpha_composite(pic3, (2222, 1008))
    d.ellipse((920, 1530, 1320, 1930), fill=rgb(GOLD) + (255,))
    d.text((970, 1640), "LIVE", font=font(DISPLAY, 76), fill=rgb(BLACK) + (255,))
    d.text((970, 1715), "SESSION", font=font(DISPLAY, 76), fill=rgb(BLACK) + (255,))
    d.ellipse((1940, 365, 2290, 715), fill=rgb(BLUE) + (255,))
    d.text((1998, 485), "ARTIST", font=font(DISPLAY, 62), fill=rgb(BLACK) + (255,))
    d.text((1998, 548), "DISCOVERY", font=font(DISPLAY, 62), fill=rgb(BLACK) + (255,))
    add_bullets(
        d,
        160,
        1980,
        [
            "Show real microphones, camera rigs, artists, producers, and performance spaces.",
            "Use cream pages for editorial storytelling and black pages for energy.",
            "Use orange stickers and blue panels as functional signals, not decoration.",
        ],
        2760,
        fnt=font(BODY, 42),
        color=INK,
        bullet_color=ORANGE_2,
        gap=14,
    )
    pages.append(save_slide(img, 8))

    # 09 Motion and Lottie
    img = gradient_bg(BLUE, "#71C7E5")
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d, dark=True)
    page_number(d, 9, dark=True)
    d.text((140, 268), "MOTION SYSTEM", font=font(DISPLAY, 168), fill=rgb(INK) + (255,))
    d.text((145, 462), "Lottie animations should feel like studio signals, not generic decoration.", font=font(BODY_BOLD, 52), fill=rgb(INK) + (230,))
    modules = [
        ("WAVEFORM REVEAL", "Hero and section headers. Audio lines draw in as the headline locks up."),
        ("ROTATING OPEN CALL", "Talent badges cycle: vocalists, session musicians, hosts, producers."),
        ("MIC PULSE", "Subtle loop for recording CTAs and booking states."),
        ("ARTIST CARD LOAD", "Profile photo, bio, socials, and video card animate in sequence."),
    ]
    x, y = 155, 710
    for i, (head, body) in enumerate(modules):
        bx = x + (i % 2) * 1440
        by = y + (i // 2) * 520
        card(d, (bx, by, bx + 1240, by + 380), CREAM_2, outline=INK, radius=18)
        d.text((bx + 48, by + 44), f"0{i+1}", font=font(MONO, 42), fill=rgb(ORANGE_2) + (255,))
        d.text((bx + 48, by + 102), head, font=font(DISPLAY, 66), fill=rgb(INK) + (255,))
        draw_wrapped(d, (bx + 48, by + 184), body, font(BODY, 36), rgb(INK) + (225,), 960, 1.18)
        draw_wave(d, bx + 750, by + 255, 360, 70, ORANGE_2)
    d.text((155, 1865), "IMPLEMENTATION RULES", font=font(DISPLAY, 74), fill=rgb(INK) + (255,))
    add_bullets(
        d,
        160,
        1975,
        [
            "Use JSON Lottie for waveform, equalizer, spinner, and badge loops.",
            "Keep loops short: 2 to 5 seconds, easing smooth, no constant full-page motion.",
            "Respect reduced-motion settings in Svelte.",
            "Never animate core text in a way that blocks reading.",
        ],
        2600,
        fnt=font(BODY, 40),
        color=INK,
        bullet_color=ORANGE_2,
        gap=16,
    )
    pages.append(save_slide(img, 9))

    # 10 Discover artists
    img = gradient_bg(CREAM, CREAM_2)
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d, dark=True)
    page_number(d, 10, dark=True)
    d.text((140, 260), "DISCOVER ARTISTS", font=font(DISPLAY, 162), fill=rgb(INK) + (255,))
    d.text((146, 445), "Every performance session becomes a discoverable profile.", font=font(BODY_BOLD, 54), fill=rgb(INK) + (235,))
    card(d, (170, 720, 1350, 1835), INK, radius=26)
    profile = cover_crop(artist, (470, 470), focus=(0.55, 0.35))
    img.alpha_composite(profile, (255, 805))
    d.rectangle((790, 805, 1265, 1072), fill=rgb(ORANGE) + (255,))
    d.text((828, 850), "FEATURED", font=font(DISPLAY, 62), fill=rgb(BLACK) + (255,))
    d.text((828, 920), "ARTIST", font=font(DISPLAY, 62), fill=rgb(BLACK) + (255,))
    d.text((255, 1320), "AMINATA K.", font=font(DISPLAY, 74), fill=rgb(CREAM) + (255,))
    draw_wrapped(d, (255, 1402), "Singer-songwriter. Live at VYBE Central Session 01.", font(BODY, 34), rgb(CREAM) + (220,), 860, 1.18)
    d.rectangle((255, 1600, 880, 1708), outline=rgb(BLUE) + (255,), width=4)
    d.text((290, 1630), "WATCH PERFORMANCE", font=font(DISPLAY, 52), fill=rgb(BLUE) + (255,))
    fields = [
        ("PROFILE PHOTO", "Consistent crop, strong expression, no heavy filters."),
        ("BIO", "Short, human, specific: genre, origin, current project."),
        ("SOCIAL LINKS", "Instagram, TikTok, YouTube, WhatsApp when relevant."),
        ("PERFORMANCE VIDEO", "Embedded session video with artist name and track title."),
    ]
    y = 745
    for i, (head, body) in enumerate(fields):
        d.text((1510, y), f"0{i+1}", font=font(MONO, 50), fill=rgb(ORANGE_2) + (255,))
        d.text((1625, y), head, font=font(DISPLAY, 72), fill=rgb(INK) + (255,))
        draw_wrapped(d, (1628, y + 86), body, font(BODY, 38), rgb(INK) + (220,), 1080, 1.18)
        y += 300
    pages.append(save_slide(img, 10))

    # 11 Community and calls
    img, d = page_base(BLACK)
    top_mark(d)
    page_number(d, 11)
    d.text((140, 260), "COMMUNITY", font=font(DISPLAY, 166), fill=rgb(CREAM) + (255,))
    d.text((140, 430), "AND OPEN CALLS", font=font(DISPLAY, 166), fill=rgb(ORANGE) + (255,))
    draw_wrapped(
        d,
        (148, 640),
        "The brand is strongest when people can join, contribute, and see a path into the creative ecosystem.",
        font(BODY_BOLD, 54),
        rgb(CREAM) + (235,),
        1280,
        1.18,
    )
    card(d, (165, 930, 1320, 1765), "#17120F", outline="#382D24", radius=24)
    d.text((235, 1000), "JOIN THE COMMUNITY", font=font(DISPLAY, 74), fill=rgb(BLUE) + (255,))
    for i, label in enumerate(["Email address", "WhatsApp number", "Creative role"]):
        by = 1135 + i * 170
        d.rounded_rectangle((235, by, 1240, by + 108), radius=10, outline=rgb(CREAM) + (120,), width=3)
        d.text((268, by + 35), label, font=font(BODY, 34), fill=rgb(CREAM) + (160,))
    d.rounded_rectangle((235, 1660, 740, 1775), radius=60, fill=rgb(ORANGE) + (255,))
    d.text((295, 1692), "JOIN NOW", font=font(DISPLAY, 58), fill=rgb(BLACK) + (255,))
    calls = ["BACKGROUND VOCALISTS", "SESSION MUSICIANS", "HOSTS", "PRODUCERS"]
    x = 1600
    y = 820
    for i, item in enumerate(calls):
        by = y + i * 292
        d.text((x, by), f"{i+1:02d}", font=font(MONO, 56), fill=rgb(BLUE) + (255,))
        d.text((x + 145, by), item, font=font(DISPLAY, 82), fill=rgb(CREAM) + (255,))
        d.line((x + 145, by + 112, x + 1260, by + 112), fill=rgb(ORANGE) + (255,), width=5)
    d.ellipse((2460, 420, 2910, 870), fill=rgb(ORANGE) + (255,))
    d.text((2552, 560), "OPEN", font=font(DISPLAY, 88), fill=rgb(BLACK) + (255,))
    d.text((2552, 645), "CALLS", font=font(DISPLAY, 88), fill=rgb(BLACK) + (255,))
    pages.append(save_slide(img, 11))

    # 12 Touchpoints
    img = gradient_bg("#EF4D24", "#F49222")
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d)
    page_number(d, 12)
    d.text((140, 255), "TOUCHPOINTS", font=font(DISPLAY, 172), fill=rgb(WHITE) + (255,))
    d.text((146, 444), "The brand needs to work from social cards to studio walls.", font=font(BODY_BOLD, 54), fill=rgb(BLACK) + (230,))
    # Social card
    card(d, (170, 740, 980, 1580), BLACK, radius=18)
    img.alpha_composite(cover_crop(hero, (810, 470), focus=(0.5, 0.4)), (170, 740))
    d.rectangle((170, 740, 980, 1210), fill=rgb(BLACK) + (70,))
    d.text((230, 1270), "LIVE SESSION", font=font(DISPLAY, 72), fill=rgb(ORANGE) + (255,))
    d.text((230, 1350), "OUT NOW", font=font(DISPLAY, 92), fill=rgb(CREAM) + (255,))
    # Merchandise block
    card(d, (1190, 740, 1995, 1580), "#171717", radius=18)
    d.rounded_rectangle((1395, 925, 1785, 1440), radius=72, fill=rgb("#242424") + (255,))
    paste_logo(img, logo, (1466, 1035, 1718, 1287), shadow=False)
    d.text((1270, 1510), "MERCH / STAFF / CREW", font=font(DISPLAY, 52), fill=rgb(CREAM) + (255,))
    # Signage
    card(d, (2200, 740, 3030, 1580), CREAM_2, radius=18)
    paste_logo(img, logo, (2370, 850, 2860, 1340), shadow=True)
    d.text((2315, 1488), "STUDIO SIGNAGE", font=font(DISPLAY, 56), fill=rgb(INK) + (255,))
    add_bullets(
        d,
        175,
        1790,
        [
            "Social: use black cards, orange labels, artist photography, short CTAs.",
            "Studio: keep physical signage bold, high contrast, and easy to recognize from distance.",
            "Merch: black base first; orange logo and cream typography.",
            "Footer: Instagram, TikTok, YouTube, WhatsApp, Email, Location.",
        ],
        2800,
        fnt=font(BODY, 42),
        color=BLACK,
        bullet_color=BLACK,
        gap=14,
    )
    pages.append(save_slide(img, 12))

    # 13 Rules
    img = gradient_bg(CREAM, CREAM_2)
    d = ImageDraw.Draw(img, "RGBA")
    top_mark(d, dark=True)
    page_number(d, 13, dark=True)
    d.text((140, 260), "USAGE RULES", font=font(DISPLAY, 168), fill=rgb(INK) + (255,))
    col1 = (160, 620, 1505, 1910)
    col2 = (1695, 620, 3040, 1910)
    card(d, col1, INK, radius=22)
    card(d, col2, "#F3DDC4", outline=INK, radius=22)
    d.text((240, 700), "DO", font=font(DISPLAY, 120), fill=rgb(BLUE) + (255,))
    add_bullets(
        d,
        245,
        885,
        [
            "Use real studio and artist imagery.",
            "Keep CTAs bright orange or blue.",
            "Protect the logo clear space.",
            "Use cream for readable editorial sections.",
            "Let motion support the rhythm of the page.",
        ],
        1080,
        fnt=font(BODY, 42),
        color=CREAM,
        bullet_color=ORANGE,
        gap=20,
    )
    d.text((1775, 700), "DON'T", font=font(DISPLAY, 120), fill=rgb(ORANGE_2) + (255,))
    add_bullets(
        d,
        1780,
        885,
        [
            "Do not use flat corporate gradients.",
            "Do not place the logo over bright clutter.",
            "Do not make every section orange.",
            "Do not use stock images that hide the subject.",
            "Do not animate text so much that it becomes hard to read.",
        ],
        1080,
        fnt=font(BODY, 42),
        color=INK,
        bullet_color=ORANGE_2,
        gap=20,
    )
    d.text((165, 2055), "LAUNCH CHECK: logo PNG, favicon, palette tokens, type imports, Lottie JSON, artist-card template, footer contacts.", font=font(BODY_BOLD, 42), fill=rgb(INK) + (230,))
    pages.append(save_slide(img, 13))

    # 14 Closing
    img, d = page_base(BLACK)
    top_mark(d)
    page_number(d, 14)
    paste_logo(img, logo, (1165, 360, 2035, 1230), shadow=True)
    d.text((600, 1395), "CREATE. COLLABORATE.", font=font(DISPLAY, 170), fill=rgb(CREAM) + (255,))
    d.text((875, 1565), "BE HEARD.", font=font(DISPLAY, 190), fill=rgb(ORANGE) + (255,))
    draw_wrapped(
        d,
        (W / 2, 1805),
        "VYBE Central is more than a studio. It is the place where talent becomes visible.",
        font(BODY_BOLD, 58),
        rgb(CREAM) + (230,),
        1550,
        1.2,
        anchor="center",
    )
    footer = "Instagram / TikTok / YouTube / WhatsApp / Email / Location"
    tw, _ = text_size(d, footer, font(MONO, 34))
    d.text(((W - tw) / 2, 2180), footer, font=font(MONO, 34), fill=rgb(BLUE) + (210,))
    pages.append(save_slide(img, 14))

    return pages, logo


def build_docx(pages):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(12.8)
    section.page_height = Inches(10)
    section.top_margin = Inches(0.20)
    section.bottom_margin = Inches(0.20)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.header_distance = Inches(0)
    section.footer_distance = Inches(0)
    for i, page in enumerate(pages):
        if i:
            doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = 0
        p.paragraph_format.space_after = 0
        run = p.add_run()
        run.add_picture(str(page), width=Inches(11.90), height=Inches(9.30))
    doc.save(OUT_DOCX)


def build_preview_pdf(pages):
    images = [Image.open(p).convert("RGB") for p in pages]
    images[0].save(OUT_PDF_DIRECT, save_all=True, append_images=images[1:], resolution=250)


def main():
    SLIDES.mkdir(parents=True, exist_ok=True)
    ASSETS.mkdir(parents=True, exist_ok=True)
    for old in SLIDES.glob("page-*.png"):
        old.unlink()
    pages, logo = build_pages()
    build_docx(pages)
    build_preview_pdf(pages)
    print(f"Created {OUT_DOCX}")
    print(f"Created {OUT_PDF_DIRECT}")
    print(f"Created {logo}")


if __name__ == "__main__":
    main()
